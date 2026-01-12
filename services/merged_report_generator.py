"""
Генератор сводных отчётов (объединение данных от разных БКИ)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any
import os
from datetime import datetime

from config.settings import settings


def generate_merged_word_report(merged_data: Dict[str, Any], tariff: str, explanation: str) -> str:
    """
    Генерирует сводный Word отчёт с данными от всех БКИ
    
    Args:
        merged_data: Сводные данные от merge_bki_reports
        tariff: Рекомендованный тариф
        explanation: Объяснение выбора тарифа
        
    Returns:
        Путь к сгенерированному файлу
    """
    doc = Document()
    summary = merged_data["summary"]
    client_name = merged_data["client_name"] or "Клиент"
    
    # === ЗАГОЛОВОК ===
    title = doc.add_heading('СВОДНЫЙ КРЕДИТНЫЙ ОТЧЁТ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(f'Клиент: {client_name}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.bold = True
    
    date_p = doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Пустая строка
    
    # === ОБЩАЯ ИНФОРМАЦИЯ ===
    doc.add_heading('📊 СВОДНАЯ ИНФОРМАЦИЯ', 1)
    
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_rows = [
        ('Количество проверенных БКИ:', str(summary['total_reports'])),
        ('Источники данных:', ', '.join(summary['bki_types'])),
        ('Средний кредитный рейтинг:', f"{summary['avg_credit_score']:.0f} баллов"),
        ('Общая долговая нагрузка:', f"{summary['total_debt']:,.0f} руб"),
        ('Активных кредитных продуктов:', str(summary['total_active_accounts'])),
        ('Максимальная просрочка:', f"{summary['max_delinquency_days']} дней" if summary['max_delinquency_days'] > 0 else "Отсутствует")
    ]
    
    for i, (label, value) in enumerate(info_rows):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        # Делаем первый столбец жирным
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # === ДАННЫЕ ПО КАЖДОМУ БКИ ===
    doc.add_heading('📋 ДЕТАЛИЗАЦИЯ ПО БКИ', 1)
    
    for bki_type, bki_info in merged_data["bki_data"].items():
        doc.add_heading(bki_type, 2)
        
        bki_table = doc.add_table(rows=4, cols=2)
        bki_table.style = 'Light List Accent 1'
        
        bki_rows = [
            ('Кредитный рейтинг:', f"{bki_info['credit_score']} баллов"),
            ('Общий долг:', f"{bki_info['total_debt']:,.0f} руб"),
            ('Активных счетов:', str(bki_info['active_accounts'])),
            ('Просрочки:', f"{bki_info['max_delinquency']} дней" if bki_info['max_delinquency'] > 0 else "Нет")
        ]
        
        for i, (label, value) in enumerate(bki_rows):
            bki_table.rows[i].cells[0].text = label
            bki_table.rows[i].cells[1].text = value
            bki_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
    
    # === ВСЕ КРЕДИТЫ ===
    if merged_data["all_credits"]:
        doc.add_heading('💳 ВСЕ КРЕДИТНЫЕ ПРОДУКТЫ', 1)
        
        credits_table = doc.add_table(rows=1, cols=5)
        credits_table.style = 'Light Grid Accent 1'
        
        # Заголовки
        headers = ['Источник', 'Кредитор', 'Тип', 'Остаток долга', 'Просрочка']
        for i, header in enumerate(headers):
            cell = credits_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Данные
        for credit in merged_data["all_credits"]:
            row = credits_table.add_row()
            row.cells[0].text = credit['bki_source']
            row.cells[1].text = credit['creditor'] or 'Не указано'
            row.cells[2].text = credit['product_type'] or 'Не указано'
            row.cells[3].text = f"{credit['balance']:,.0f} руб"
            row.cells[4].text = f"{credit['delinquency_days']} дней" if credit['delinquency_days'] > 0 else "Нет"
        
        doc.add_paragraph()
    
    # === РЕКОМЕНДАЦИЯ ===
    doc.add_heading('✅ РЕКОМЕНДАЦИЯ', 1)
    
    # Рамка с тарифом
    tariff_p = doc.add_paragraph()
    tariff_run = tariff_p.add_run(f'РЕКОМЕНДУЕМЫЙ ТАРИФ: {tariff.upper()}')
    tariff_run.font.size = Pt(16)
    tariff_run.font.bold = True
    if tariff == "Premium":
        tariff_run.font.color.rgb = RGBColor(0, 128, 0)  # Зелёный
    else:
        tariff_run.font.color.rgb = RGBColor(255, 140, 0)  # Оранжевый
    tariff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Объяснение
    doc.add_heading('Обоснование:', 2)
    explanation_p = doc.add_paragraph(explanation)
    explanation_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # === ПОДПИСЬ ===
    doc.add_paragraph('_' * 50)
    signature_p = doc.add_paragraph()
    signature_p.add_run('Автоматически сгенерировано системой CreditAI\n').font.italic = True
    signature_p.add_run(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}').font.italic = True
    signature_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # === СОХРАНЕНИЕ ===
    os.makedirs(settings.OUTPUT_FOLDER, exist_ok=True)
    
    filename = f"Сводный_Отчёт_{client_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(settings.OUTPUT_FOLDER, filename)
    
    doc.save(filepath)
    
    return filepath



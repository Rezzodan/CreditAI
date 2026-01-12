"""
Генератор отчётов Word по техническому заданию
Создаёт документ с анализом кредитной истории и рекомендациями тарифов
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List
from pathlib import Path


class TZReportGenerator:
    """Генератор отчётов по ТЗ в формате Word"""
    
    def __init__(self, output_folder: str = "./output"):
        self.output_folder = Path(output_folder)
        self.output_folder.mkdir(exist_ok=True)
    
    def generate_report(self, analysis: Dict, client_name: str = "ФИО клиента") -> str:
        """
        Генерация Word отчёта по ТЗ
        
        Args:
            analysis: Результат анализа от ReportAnalyzer
            client_name: ФИО клиента
            
        Returns:
            Путь к сгенерированному файлу
        """
        doc = Document()
        
        # Настройка стилей
        self._setup_styles(doc)
        
        # Заголовок отчёта
        self._add_header(doc, analysis, client_name)
        
        # Раздел 1: Кредитный рейтинг
        self._add_rating_section(doc, analysis['sections']['rating'])
        
        # Раздел 2: Признак банкротства (если есть)
        if 'bankruptcy' in analysis['sections']:
            self._add_bankruptcy_section(doc, analysis['sections']['bankruptcy'])
        
        # Раздел 3: Текущие кредиты
        self._add_active_credits_section(doc, analysis['sections']['active_credits'])
        
        # Раздел 4: Закрытые кредиты
        self._add_closed_credits_section(doc, analysis['sections']['closed_credits'])
        
        # Раздел 5: Кредитная нагрузка
        self._add_credit_load_section(doc, analysis['sections']['credit_load'])
        
        # Раздел 6: Рекомендации
        self._add_recommendations_section(doc, analysis['sections']['recommendations'])
        
        # Сохранение
        filename = f"Отчет_{analysis['bki_type']}_{client_name}_{analysis['report_date']}.docx"
        filepath = self.output_folder / filename
        doc.save(str(filepath))
        
        return str(filepath)
    
    def _setup_styles(self, doc: Document):
        """Настройка стилей документа"""
        # Стиль для заголовков
        styles = doc.styles
        
        # Можно добавить кастомные стили
        # Пока используем стандартные
        pass
    
    def _add_header(self, doc: Document, analysis: Dict, client_name: str):
        """Добавление заголовка отчёта"""
        # Главный заголовок
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f'Результат анализа кредитной истории «{client_name}»')
        run.font.size = Pt(16)
        run.font.bold = True
        
        # Подзаголовок
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('на основании предоставленных кредитных отчетов')
        run.font.size = Pt(12)
        
        # Дата и тип БКИ
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(
            f"\nРезультат анализа кредитного отчета {analysis['bki_type']} "
            f"от «{analysis['report_date']}»"
        )
        run.font.size = Pt(12)
        run.font.bold = True
        
        doc.add_paragraph()  # Пустая строка
    
    def _add_rating_section(self, doc: Document, rating_data: Dict):
        """Раздел 1: Кредитный рейтинг"""
        heading = doc.add_heading('Раздел 1. Текущий кредитный рейтинг', level=2)
        
        p = doc.add_paragraph()
        run = p.add_run(f"Ваш кредитный рейтинг: ")
        run.font.size = Pt(12)
        
        run = p.add_run(f"{rating_data['value']}")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 112, 192)  # Синий цвет
        
        if rating_data.get('has_image'):
            p = doc.add_paragraph()
            p.add_run("(Место для картинки рейтинга из отчёта)")
            p.italic = True
        
        doc.add_paragraph()
    
    def _add_bankruptcy_section(self, doc: Document, bankruptcy_data: Dict):
        """Раздел 2: Признак банкротства"""
        heading = doc.add_heading('Раздел 2. Признак банкротства', level=2)
        
        p = doc.add_paragraph()
        status = "Да" if bankruptcy_data['has_bankruptcy'] else "Нет"
        run = p.add_run(f"Признак банкротства: {status}")
        run.font.size = Pt(12)
        
        if bankruptcy_data['has_bankruptcy']:
            run.font.color.rgb = RGBColor(192, 0, 0)  # Красный
            run.font.bold = True
        
        doc.add_paragraph()
    
    def _add_active_credits_section(self, doc: Document, credits_data: Dict):
        """Раздел: Текущие кредиты"""
        section_num = '3' if 'bankruptcy' not in credits_data else '3'
        heading = doc.add_heading(f'Раздел {section_num}. Текущие кредиты', level=2)
        
        if not credits_data['has_credits']:
            p = doc.add_paragraph("Текущие кредиты отсутствуют")
            p.italic = True
            doc.add_paragraph()
            return
        
        p = doc.add_paragraph()
        run = p.add_run(f"Количество активных кредитов: {len(credits_data['credits'])}")
        run.font.bold = True
        
        # Детали по каждому кредиту
        for i, credit in enumerate(credits_data['credits'], 1):
            doc.add_paragraph()
            
            # Заголовок кредита
            p = doc.add_paragraph()
            run = p.add_run(f"Кредит {i}:")
            run.font.bold = True
            run.font.size = Pt(12)
            
            # Детали
            details = [
                f"Наименование кредитора: {credit['creditor']}",
                f"Вид кредита: {credit['type']}",
                f"Сумма: {credit['amount']:,.2f} руб.",
                f"Дата открытия: {credit['open_date']}",
                f"Ежемесячный платеж: {credit['payment']:,.2f} руб.",
                f"Остаток задолженности: {credit['balance']:,.2f} руб.",
            ]
            
            for detail in details:
                doc.add_paragraph(detail, style='List Bullet')
            
            # Информация о просрочках
            if credit['has_overdue']:
                p = doc.add_paragraph()
                run = p.add_run("⚠ Наличие просрочек: Да")
                run.font.color.rgb = RGBColor(255, 128, 0)  # Оранжевый
                run.font.bold = True
                
                doc.add_paragraph(
                    f"Максимальная длительность просрочек: {credit['max_overdue_days']} дней",
                    style='List Bullet'
                )
            else:
                p = doc.add_paragraph()
                run = p.add_run("✓ Наличие просрочек: Нет")
                run.font.color.rgb = RGBColor(0, 176, 80)  # Зелёный
            
            # Текущая просрочка
            if credit['current_overdue']:
                p = doc.add_paragraph()
                run = p.add_run("⚠ ТЕКУЩАЯ ПРОСРОЧКА!")
                run.font.color.rgb = RGBColor(192, 0, 0)  # Красный
                run.font.bold = True
                
                doc.add_paragraph(
                    f"Количество дней: {credit['current_overdue_days']} дней",
                    style='List Bullet'
                )
                doc.add_paragraph(
                    f"Сумма: {credit['current_overdue_amount']:,.2f} руб.",
                    style='List Bullet'
                )
        
        doc.add_paragraph()
    
    def _add_closed_credits_section(self, doc: Document, credits_data: Dict):
        """Раздел: Закрытые кредиты"""
        heading = doc.add_heading('Раздел 4. Закрытые кредиты', level=2)
        
        if not credits_data['has_credits']:
            p = doc.add_paragraph("Закрытые кредиты отсутствуют")
            p.italic = True
            doc.add_paragraph()
            return
        
        p = doc.add_paragraph()
        run = p.add_run(f"Количество закрытых кредитов: {len(credits_data['credits'])}")
        run.font.bold = True
        
        # Детали по каждому закрытому кредиту
        for i, credit in enumerate(credits_data['credits'], 1):
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            run = p.add_run(f"Кредит {i}:")
            run.font.bold = True
            
            details = [
                f"Наименование кредитора: {credit['creditor']}",
                f"Вид кредита: {credit['type']}",
                f"Сумма: {credit['amount']:,.2f} руб.",
                f"Дата открытия: {credit['open_date']}",
            ]
            
            for detail in details:
                doc.add_paragraph(detail, style='List Bullet')
            
            if credit['has_overdue']:
                p = doc.add_paragraph()
                run = p.add_run("⚠ Наличие просрочек: Да")
                run.font.color.rgb = RGBColor(255, 128, 0)
                run.font.bold = True
                
                doc.add_paragraph(
                    f"Максимальная длительность: {credit['max_overdue_days']} дней",
                    style='List Bullet'
                )
            else:
                p = doc.add_paragraph()
                run = p.add_run("✓ Наличие просрочек: Нет")
                run.font.color.rgb = RGBColor(0, 176, 80)
        
        doc.add_paragraph()
    
    def _add_credit_load_section(self, doc: Document, load_data: Dict):
        """Раздел: Кредитная нагрузка"""
        heading = doc.add_heading('Раздел 5. Кредитная нагрузка', level=2)
        
        # Текущая задолженность
        p = doc.add_paragraph()
        p.add_run("Текущая кредитная задолженность: ").font.bold = True
        run = p.add_run(f"{load_data['current_debt']:,.2f} руб.")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 112, 192)
        
        # Просроченная задолженность
        p = doc.add_paragraph()
        p.add_run("Просроченная кредитная задолженность: ").font.bold = True
        run = p.add_run(f"{load_data['overdue_debt']:,.2f} руб.")
        run.font.size = Pt(12)
        
        if load_data['overdue_debt'] > 0:
            run.font.color.rgb = RGBColor(192, 0, 0)  # Красный если есть
            run.font.bold = True
        else:
            run.font.color.rgb = RGBColor(0, 176, 80)  # Зелёный если нет
        
        # Ежемесячный платёж (если есть)
        if 'monthly_payment' in load_data:
            p = doc.add_paragraph()
            p.add_run("Текущая кредитная нагрузка (платёж в месяц): ").font.bold = True
            run = p.add_run(f"{load_data['monthly_payment']:,.2f} руб.")
            run.font.size = Pt(12)
        
        doc.add_paragraph()
    
    def _add_recommendations_section(self, doc: Document, rec_data: Dict):
        """Раздел: Рекомендации"""
        heading = doc.add_heading('Раздел 6. Рекомендации/примечания/отклонения', level=2)
        
        # Рекомендуемый тариф
        p = doc.add_paragraph()
        p.add_run("РЕКОМЕНДУЕМЫЙ ТАРИФ: ").font.bold = True
        run = p.add_run(rec_data['tariff_name'])
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 112, 192)
        
        doc.add_paragraph()
        
        # Детальные рекомендации
        for item in rec_data['items']:
            p = doc.add_paragraph()
            
            # Иконка в зависимости от типа
            if item['type'] == 'critical':
                icon = "🔴 "
                color = RGBColor(192, 0, 0)
            elif item['type'] == 'warning':
                icon = "⚠️ "
                color = RGBColor(255, 128, 0)
            elif item['type'] == 'success':
                icon = "✓ "
                color = RGBColor(0, 176, 80)
            else:
                icon = "ℹ️ "
                color = RGBColor(0, 112, 192)
            
            run = p.add_run(icon)
            run = p.add_run(item['text'])
            run.font.size = Pt(11)
            
            # Выделяем критичные моменты
            if item['type'] in ['critical', 'warning']:
                run.font.bold = True
            
            doc.add_paragraph()
    
    def generate_combined_report(self, analyses: List[Dict], client_name: str) -> str:
        """
        Генерация объединённого отчёта из нескольких БКИ
        
        Args:
            analyses: Список анализов от разных БКИ
            client_name: ФИО клиента
            
        Returns:
            Путь к сгенерированному файлу
        """
        doc = Document()
        self._setup_styles(doc)
        
        # Общий заголовок
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f'Результат анализа кредитной истории «{client_name}»')
        run.font.size = Pt(16)
        run.font.bold = True
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('на основании предоставленных кредитных отчетов')
        run.font.size = Pt(12)
        
        doc.add_paragraph()
        doc.add_page_break()
        
        # Добавляем каждый БКИ отчёт
        for i, analysis in enumerate(analyses):
            if i > 0:
                doc.add_page_break()
            
            # Заголовок БКИ
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = info.add_run(
                f"\nРезультат анализа кредитного отчета {analysis['bki_type']} "
                f"от «{analysis['report_date']}»"
            )
            run.font.size = Pt(14)
            run.font.bold = True
            
            doc.add_paragraph()
            
            # Разделы
            self._add_rating_section(doc, analysis['sections']['rating'])
            
            if 'bankruptcy' in analysis['sections']:
                self._add_bankruptcy_section(doc, analysis['sections']['bankruptcy'])
            
            self._add_active_credits_section(doc, analysis['sections']['active_credits'])
            self._add_closed_credits_section(doc, analysis['sections']['closed_credits'])
            self._add_credit_load_section(doc, analysis['sections']['credit_load'])
            self._add_recommendations_section(doc, analysis['sections']['recommendations'])
        
        # Сохранение
        filename = f"Отчет_Комплексный_{client_name}.docx"
        filepath = self.output_folder / filename
        doc.save(str(filepath))
        
        return str(filepath)




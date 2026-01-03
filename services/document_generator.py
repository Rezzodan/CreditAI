"""
Модуль генерации документов Word
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Optional
from pathlib import Path
from datetime import datetime
import json


class DocumentGenerator:
    """Генератор документов Word"""
    
    def __init__(self, output_folder: str = "./output"):
        self.output_folder = Path(output_folder)
        self.output_folder.mkdir(parents=True, exist_ok=True)
    
    def generate_analysis_report(self, data: Dict, report_id: str) -> str:
        """
        Генерирует аналитический отчёт
        
        Args:
            data: Извлечённые данные
            report_id: ID отчёта
            
        Returns:
            Путь к созданному файлу
        """
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Заголовок
        title = doc.add_heading('Аналитический отчёт по кредитной истории', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Метаданные
        metadata = data.get('metadata', {})
        doc.add_paragraph(f"Дата обработки: {metadata.get('processed_at', datetime.now().isoformat())}")
        doc.add_paragraph(f"Тип БКИ: {metadata.get('bki_type', 'Неизвестно')}")
        doc.add_paragraph(f"Уверенность обработки: {metadata.get('confidence_overall', 0):.2%}")
        doc.add_paragraph("")
        
        # Данные субъекта
        doc.add_heading('Данные субъекта кредитной истории', level=1)
        subject = data.get('subject', {})
        
        if 'full_name' in subject:
            doc.add_paragraph(f"ФИО: {subject['full_name'].get('value', 'Не указано')}")
        
        if 'birth_date' in subject:
            doc.add_paragraph(f"Дата рождения: {subject['birth_date'].get('value', 'Не указано')}")
        
        if 'passport' in subject:
            passport = subject['passport']
            doc.add_paragraph(
                f"Паспорт: {passport.get('series', '')} {passport.get('number', '')}"
            )
        
        doc.add_paragraph("")
        
        # Сводка
        doc.add_heading('Сводная информация', level=1)
        summary = data.get('summary', {})
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        rows_data = [
            ('Общая сумма долга', f"{summary.get('total_debt', 0):,.2f} руб."),
            ('Активных счетов', str(summary.get('active_accounts', 0))),
            ('Максимальная просрочка', f"{summary.get('max_delinquency_days', 0)} дней"),
            ('Кредитный балл', str(summary.get('credit_score', 'Не указан'))),
            ('Уровень риска', summary.get('risk_score', 'Не определён'))
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
        
        doc.add_paragraph("")
        
        # Кредитные счета
        accounts = data.get('accounts', [])
        if accounts:
            doc.add_heading('Кредитные обязательства', level=1)
            
            accounts_table = doc.add_table(rows=1, cols=7)
            accounts_table.style = 'Light Grid Accent 1'
            
            # Заголовки
            headers = ['Кредитор', 'Тип продукта', 'Лимит', 'Остаток', 'Статус', 'Просрочка', 'Дата открытия']
            header_cells = accounts_table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Данные
            for account in accounts:
                row = accounts_table.add_row()
                creditor = account.get('creditor', {})
                amounts = account.get('amounts', {})
                status = account.get('status', {})
                dates = account.get('dates', {})
                
                row.cells[0].text = creditor.get('name', 'Не указано')
                row.cells[1].text = account.get('product_type', 'Не указано')
                row.cells[2].text = f"{amounts.get('limit', 0):,.2f}" if amounts.get('limit') else "—"
                row.cells[3].text = f"{amounts.get('current_balance', 0):,.2f}" if amounts.get('current_balance') else "—"
                row.cells[4].text = status.get('general', 'Не указано')
                row.cells[5].text = f"{status.get('delinquency_days', 0)} дн." if status.get('delinquency_days') else "—"
                row.cells[6].text = dates.get('open', 'Не указано')
        
        # Рекомендации (если есть)
        if 'recommendations' in data:
            doc.add_paragraph("")
            doc.add_heading('Рекомендации', level=1)
            recommendations = data.get('recommendations', {})
            
            if 'risk_level' in recommendations:
                risk_level = recommendations['risk_level']
                doc.add_paragraph(f"Уровень риска: {risk_level}", style='List Bullet')
            
            if 'recommendations' in recommendations:
                doc.add_paragraph("Рекомендации:")
                for rec in recommendations['recommendations']:
                    doc.add_paragraph(
                        f"• {rec.get('action', '')} - {rec.get('reason', '')}",
                        style='List Bullet'
                    )
        
        # Сохранение
        filename = f"analysis_report_{report_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = self.output_folder / filename
        doc.save(str(file_path))
        
        return str(file_path)
    
    def generate_bki_letter(self, bki_type: str, client_data: Dict, 
                          errors: List[Dict], report_id: str) -> str:
        """
        Генерирует письмо в БКИ
        
        Args:
            bki_type: Тип БКИ
            client_data: Данные клиента
            errors: Список ошибок
            report_id: ID отчёта
            
        Returns:
            Путь к созданному файлу
        """
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Заголовок
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        title.add_run(f"В {bki_type}").bold = True
        
        doc.add_paragraph("")
        
        # Данные отправителя (можно настроить)
        doc.add_paragraph("От:")
        doc.add_paragraph(client_data.get('full_name', 'Не указано'))
        doc.add_paragraph(f"Паспорт: {client_data.get('passport', {}).get('series', '')} "
                         f"{client_data.get('passport', {}).get('number', '')}")
        doc.add_paragraph("")
        
        # Текст письма
        doc.add_paragraph("Уважаемые коллеги!")
        doc.add_paragraph("")
        
        doc.add_paragraph(
            "Настоящим письмом обращаюсь к Вам с просьбой об исправлении "
            "ошибок в моей кредитной истории, обнаруженных при анализе отчёта."
        )
        doc.add_paragraph("")
        
        # Список ошибок
        if errors:
            doc.add_paragraph("Обнаруженные ошибки:", style='List Bullet')
            doc.add_paragraph("")
            
            for i, error in enumerate(errors, 1):
                doc.add_paragraph(f"{i}. {error.get('field', 'Поле не указано')}", 
                                style='List Bullet')
                doc.add_paragraph(f"   Некорректное значение: {error.get('incorrect_value', 'Не указано')}")
                doc.add_paragraph(f"   Корректное значение: {error.get('correct_value', 'Не указано')}")
                if error.get('evidence'):
                    doc.add_paragraph(f"   Обоснование: {error.get('evidence', '')}")
                doc.add_paragraph("")
        
        # Требования
        doc.add_paragraph(
            "В соответствии с Федеральным законом №218-ФЗ «О кредитных историях», "
            "прошу Вас в установленные законом сроки исправить указанные ошибки "
            "и уведомить меня о результатах рассмотрения."
        )
        doc.add_paragraph("")
        
        # Подпись
        doc.add_paragraph("")
        doc.add_paragraph("С уважением,")
        doc.add_paragraph(client_data.get('full_name', ''))
        doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
        
        # Сохранение
        filename = f"bki_letter_{bki_type}_{report_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = self.output_folder / filename
        doc.save(str(file_path))
        
        return str(file_path)
    
    def generate_summary(self, data: Dict, report_id: str) -> str:
        """
        Генерирует краткую сводку
        
        Args:
            data: Извлечённые данные
            report_id: ID отчёта
            
        Returns:
            Путь к созданному файлу
        """
        doc = Document()
        
        # Заголовок
        doc.add_heading('Краткая сводка по кредитной истории', 0)
        
        summary = data.get('summary', {})
        subject = data.get('subject', {})
        
        doc.add_paragraph(f"Клиент: {subject.get('full_name', {}).get('value', 'Не указано')}")
        doc.add_paragraph(f"Общая сумма долга: {summary.get('total_debt', 0):,.2f} руб.")
        doc.add_paragraph(f"Активных счетов: {summary.get('active_accounts', 0)}")
        doc.add_paragraph(f"Уровень риска: {summary.get('risk_score', 'Не определён')}")
        
        # Сохранение
        filename = f"summary_{report_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = self.output_folder / filename
        doc.save(str(file_path))
        
        return str(file_path)



